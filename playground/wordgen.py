from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_heading(doc, text, level=1):
    """Add a heading with specified level and Japanese font"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'MS Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')


def add_japanese_paragraph(doc, text, indent=False):
    """Add a paragraph with Japanese font"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'MS Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
    return p


def insert_equation(doc, equation):
    """Insert an equation into the document"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    equation_element = OxmlElement('m:oMath')
    equation_element.text = equation
    r._element.append(equation_element)


def create_report():
    # Create document
    doc = Document()

    # Title
    create_heading(doc, "卒論の図の訂正と変換効率の数値解の計算", 0)
    add_japanese_paragraph(doc, "名前: 辛島実")
    doc.add_paragraph()  # Add spacing

    # Background and Purpose
    create_heading(doc, "背景と目的")
    add_japanese_paragraph(doc, "• 卒論の一部で使用していた先行研究の図面の修正が必要")
    add_japanese_paragraph(
        doc, "• 不均一な分極反転周期における変換効率の算出のため、従来の近似解析解に代えて非線形モード結合方程式による数値解を導出")
    doc.add_paragraph()

    # What was done
    create_heading(doc, "やったこと")

    # 1. Thesis figure modification
    create_heading(doc, "1. 卒論図面の修正", 2)
    add_japanese_paragraph(doc, "4種類の図面を新たに作成しました。")
    add_japanese_paragraph(doc, "（図1-4: 添付資料参照）")
    doc.add_paragraph()

    # 2. Calculation and verification of conversion efficiency
    create_heading(doc, "2. 変換効率の計算と検証", 2)

    # (1) Comparison with existing research
    create_heading(doc, "(1) 既存研究との比較検証", 3)
    add_japanese_paragraph(doc, "• 野呂氏による計算結果")
    add_japanese_paragraph(doc, "（図5: 添付資料参照）")

    add_japanese_paragraph(doc, "• Pythonによる解析解の再現計算")
    add_japanese_paragraph(doc, "（図6: 添付資料参照）")

    add_japanese_paragraph(doc, "• 非線形モード結合方程式による数値解と解析解の比較")
    add_japanese_paragraph(doc, "（図7: 添付資料参照）")
    doc.add_paragraph()

    # (2) Analysis of period fluctuation effects
    create_heading(doc, "(2) 周期変動による影響の解析", 3)
    add_japanese_paragraph(
        doc, "Point index（z軸方向の微小区間index）とdelta value（2Δの値）の関係を検証")
    add_japanese_paragraph(doc, "• 野呂氏の条件における2Δの値を青線で表示")
    add_japanese_paragraph(doc, "（図8,9: 添付資料参照）")
    doc.add_paragraph()

    # Equation
    add_japanese_paragraph(doc, "2Δの理論式：")
    # Note: This is a placeholder for the equation. In practice, you might need to use a different method
    # to insert the actual equation, as complex equations might require additional libraries or manual formatting
    add_japanese_paragraph(doc, "2Δ^(q)_SHG = β^(2ω) - (2β^ω + qK)")
    doc.add_paragraph()

    # Summary
    create_heading(doc, "まとめ")
    add_japanese_paragraph(doc, "1. 卒論における計算過程の検証と図面の刷新を行った")
    add_japanese_paragraph(doc, "2. 野呂氏の導波路モデルにおいて、数値解と解析解が近似していることを確認")
    doc.add_paragraph()

    # Attachments
    create_heading(doc, "添付資料")
    add_japanese_paragraph(doc, "• 図1-4: 修正した卒論図面")
    add_japanese_paragraph(doc, "• 図5: 野呂氏の計算結果")
    add_japanese_paragraph(doc, "• 図6: Python による解析解の計算結果")
    add_japanese_paragraph(doc, "• 図7: 数値解と解析解の比較")
    add_japanese_paragraph(doc, "• 図8-9: 周期変動の影響解析結果")

    # Save the document
    doc.save('thesis_report.docx')


if __name__ == "__main__":
    create_report()
